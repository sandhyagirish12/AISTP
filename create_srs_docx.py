#!/usr/bin/env python3
"""
Generate SRS.docx from SRS.md using markdown to word conversion
"""
import subprocess
import sys
import os

# Try installing python-docx if not available
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('Software Requirements Specification (SRS)')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('AI Safety Toolkit for Open-Weight Model Outputs')
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    doc.add_paragraph('Document Version: 1.0')
    doc.add_paragraph('Date: July 9, 2026')
    doc.add_paragraph('Project Status: Active Development')
    doc.add_paragraph('')
    
    # Read and parse markdown content
    with open('SRS.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_level = 0
    for line in lines:
        line = line.rstrip('\n')
        if not line.strip():
            doc.add_paragraph('')
            continue
        
        # Heading parsing
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('✓ '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('[ ] '):
            doc.add_paragraph(line[4:], style='List Bullet')
        elif line.startswith('| '):
            # Skip table markers - will be handled separately
            pass
        else:
            # Regular paragraph
            if line.strip():
                doc.add_paragraph(line)
    
    doc.save('SRS.docx')
    print('[+] SRS.docx created successfully!')
    print(f'  Location: {os.path.abspath("SRS.docx")}')
    print(f'  Size: {os.path.getsize("SRS.docx") / 1024:.1f} KB')

if __name__ == '__main__':
    create_docx()
